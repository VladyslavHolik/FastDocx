#!/usr/bin/env python
# -*- coding: utf8 -*-

import os
from docx.shared import Pt
from docx import Document
import sys
from PySide2.QtCore import (QCoreApplication, QDate, QDateTime, QMetaObject,
    QObject, QPoint, QRect, QSize, QTime, QUrl, Qt)
from PySide2.QtGui import (QBrush, QColor, QConicalGradient, QCursor, QFont,
    QFontDatabase, QIcon, QKeySequence, QLinearGradient, QPalette, QPainter,
    QPixmap, QRadialGradient)
from PySide2.QtWidgets import *
import PySide2
from PySide2.QtWidgets import QMessageBox
from design1 import Ui_Form
import pickle
import datetime
from docx.shared import Cm

def resource_path(relative):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, relative)
    return os.path.join(relative)

def add_template():
    template = ui.add_pattern_edit.text()
    ui.add_pattern_edit.clear()
    ui.comboBox.addItem(template)
    list_of_templ.append(template)
    index = ui.comboBox.findText(template, PySide2.QtCore.Qt.MatchFixedString)
    ui.comboBox.setCurrentIndex(index)
    with open('./data/templates.dat', "wb") as f:
        pickle.dump(list_of_templ, f)

def remove_template():
    what_to_remove = ui.comboBox.currentText()
    list_of_templ.remove(what_to_remove)
    index = ui.comboBox.findText(what_to_remove, PySide2.QtCore.Qt.MatchFixedString)
    ui.comboBox.removeItem(index)
    with open('./data/templates.dat', "wb") as f:
        pickle.dump(list_of_templ, f)

def show_popup():
    msg = QMessageBox()
    msg.setWindowIcon(QIcon("./img/splash.ico"))
    msg.setWindowTitle("Недійсні дані")
    msg.setText("Не обраний жодний документ огляду!")
    msg.setIcon(QMessageBox.Question)
    x = msg.exec_()

def show_necessary():
    msg = QMessageBox()
    msg.setWindowIcon(QIcon("./img/splash.ico"))
    msg.setWindowTitle("Недійсні дані")
    msg.setText("Не введені деякі дані!")
    msg.setIcon(QMessageBox.Question)
    x = msg.exec_()

def show_greet():
    msg = QMessageBox()
    msg.setWindowIcon(QIcon("./img/splash.ico"))
    msg.setWindowTitle("Завершено")
    if ui.checkBox.isChecked() and ui.checkBox_2.isChecked():
         msg.setText("Файли створені")
    else:
        msg.setText("Файл створений")
    msg.setIcon(QMessageBox.Question)
    x = msg.exec_()

def selectfolder():
    global directory
    directory = str(QFileDialog.getExistingDirectory(Dialog, "Select Directory", "C:\\"))
    ui.folder.setText(directory)

def createDocument():
    """Function for creating document for printing."""

    if not (ui.checkBox.isChecked() or ui.checkBox_2.isChecked()):
        show_popup()
        return

    name = ui.lineEdit_2.text()
    surname = ui.lineEdit.text()
    father = ui.lineEdit_3.text()
    place = ui.lineEdit_4.text()
    job = ui.comboBox.currentText()
    birthday = ui.dateEdit.text()
    date = ui.dateEdit_2.text()
    valid = ui.dateEdit_3.text()

    if not (name and surname and place and job and birthday and \
    date and valid):
        show_necessary()
        return

    if ui.checkBox.isChecked():
        filename = directory + "\Псих_огляд.docx"
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run('                                          УКРАЇНА')
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                            Київська область, м. Буча")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     КНП «Бучанський консультативно-діагностичний")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     центр» Бучанської міської ради")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     ідентифікаційний номер")
        run = paragraph.add_run(" 42081679")
        run.font.bold = True
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  \n")
        paragraph.add_run("                                                          " + surname + "\n").font.bold = True
        paragraph.add_run("                                                          " + name + "\n").font.bold = True
        if father:
            paragraph.add_run("                                                          " + father + "\n").font.bold = True
        else:
            paragraph.add_run("\n")
        paragraph.add_run("                                                             " + birthday + " року\n")
        run = paragraph.add_run(" \n")
        run.font.size = Pt(3)
        paragraph.add_run("                                  " + place + "\n")
        paragraph.add_run("                                                                                 " + date+" року")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                      " + job)
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  \n")
        paragraph.add_run("                                                                        " + valid + " року\n\n")
        run = paragraph.add_run("                                                 Голік М.Г")
        run.font.italic = True
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run(" ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                                                                                                            "+datetime.datetime.now().strftime("%d.%m.%Y")+" року")
        document.save(filename)

    # Second file

    if ui.checkBox_2.isChecked():
        filename = directory + "\Нарк_огляд.docx"
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run('                                           УКРАЇНА')
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                            Київська область, м. Буча")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     КНП «Бучанський консультативно-діагностичний")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     центр» Бучанської міської ради")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                     ідентифікаційний номер")
        run = paragraph.add_run(" 42081679")
        run.font.bold = True
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.style = document.styles['Normal']
        if not father:
            father = " "
        paragraph.add_run("                                                                        " + surname + " " + name + " " + father).font.bold = True
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                                                                 " + birthday + " року")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run("                                      " + place)
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run("                                                                     " + date+" року")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run(" ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run("                        " + job)
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("  ")
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                                                               " + valid + " року\n")
        run = paragraph.add_run("                                              Голік М.Г")
        run.font.italic = True
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        paragraph.add_run("                                                                                                    "+datetime.datetime.now().strftime("%d.%m.%Y")+" року")

        document.save(filename)

    with open("./data/directory.dat", "wb") as f:
        pickle.dump(directory, f)

    show_greet()

list_of_templ = []
with open("./data/templates.dat", "rb") as f:
    list_of_templ = pickle.load(f)

directory = ""
with open("./data/directory.dat", "rb") as dir:
    directory = pickle.load(dir)

app = PySide2.QtWidgets.QApplication(sys.argv)

splash_img = QPixmap("./img/splash.png")
splash = QSplashScreen(splash_img)
splash.show()

Dialog = PySide2.QtWidgets.QDialog()
Dialog.setWindowIcon(QIcon("./img/splash.ico"))
ui = Ui_Form()
ui.setupUi(Dialog)

for template in list_of_templ:
    ui.comboBox.addItem(template)

ui.folder.setText(directory)

splash.finish(Dialog)
Dialog.show()

ui.get_folder.clicked.connect(selectfolder)
ui.pushButton.clicked.connect(createDocument)
ui.pushButton_2.clicked.connect(add_template)
ui.deleteButton.clicked.connect(remove_template)

sys.exit(app.exec_())
